#!/usr/bin/env python3
"""
Impact Analysis Agent - Production Ready Version
Analyzes GitHub repositories and generates tech stack recommendations using OpenAI API
"""

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import json
import os
import urllib.request
import subprocess
import tempfile
import shutil
from datetime import datetime
from io import BytesIO
import uvicorn
import socket
import time
import random

# Load environment variables from .env file
def load_env():
    try:
        env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
        with open(env_path, 'r') as f:
            for line in f:
                line = line.strip()
                if line and '=' in line and not line.startswith('#'):
                    key, value = line.split('=', 1)
                    os.environ[key] = value
    except FileNotFoundError:
        pass
    except Exception as e:
        pass

load_env()

try:
    import PyPDF2
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    from docx import Document
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False

# Configuration
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
MODEL = os.getenv("MODEL", "gpt-3.5-turbo")
TEMPERATURE = float(os.getenv("TEMPERATURE", "0.3"))
API_URL = os.getenv("API_URL", "https://openrouter.ai/api/v1/chat/completions")
GENERATED_FILES_DIR = os.getenv("GENERATED_FILES_DIR", "generated_files")

def find_available_port():
    """Find an available port"""
    start_port = int(os.getenv("START_PORT", "8090"))
    max_attempts = int(os.getenv("PORT_SEARCH_RANGE", "50"))
    
    for port in range(start_port, start_port + max_attempts):
        try:
            bind_host = os.getenv("BIND_HOST", "127.0.0.1")
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.bind((bind_host, port))
                return port
        except OSError:
            continue
    
    try:
        bind_host = os.getenv("BIND_HOST", "127.0.0.1")
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.bind((bind_host, 0))
            return s.getsockname()[1]
    except OSError:
        fallback_port = int(os.getenv("FALLBACK_PORT", "8090"))
        return fallback_port

# FastAPI app
app = FastAPI(
    title="Impact Analysis Agent", 
    description="Analyzes GitHub repositories and generates tech stack recommendations",
    docs_url="/api/docs", 
    redoc_url="/api/redoc"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models
class AnalysisRequest(BaseModel):
    repo_url: str
    architecture_content: str
    prd_content: str = ""

class FileUploadResponse(BaseModel):
    success: bool
    extracted_text: str = ""
    filename: str = ""
    error: str = ""

class AnalysisResponse(BaseModel):
    success: bool
    analysis: str = ""
    document_id: str = ""
    timestamp: str = ""
    error: str = ""

# Agent class
class ImpactAnalysisAgent:
    def __init__(self):
        self.api_key = OPENROUTER_API_KEY
        self.model = MODEL
        self.temperature = TEMPERATURE
        self.api_url = API_URL

    def extract_text_from_file(self, file_data, filename):
        """Extract text from various file formats"""
        try:
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext == '.pdf' and PDF_SUPPORT:
                return self._extract_pdf_text(file_data)
            elif file_ext in ['.doc', '.docx'] and DOC_SUPPORT:
                return self._extract_doc_text(file_data)
            else:
                return file_data.decode('utf-8', errors='ignore')
        except Exception as e:
            raise Exception(f"Failed to extract text from {filename}: {str(e)}")

    def _extract_pdf_text(self, file_data):
        """Extract text from PDF files"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_data))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception("PDF extraction failed")

    def _extract_doc_text(self, file_data):
        """Extract text from DOC/DOCX files"""
        try:
            doc = docx.Document(BytesIO(file_data))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception("DOC extraction failed")

    def generate_pdf(self, content, title="Document"):
        """Generate PDF from text content"""
        if not PDF_SUPPORT:
            raise Exception("PDF generation not supported")
        
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            story.append(Paragraph(title, styles['Title']))
            story.append(Spacer(1, 12))
            
            for line in content.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            raise Exception("PDF generation failed")

    def generate_docx(self, content, title="Document"):
        """Generate DOCX from text content"""
        if not DOC_SUPPORT:
            raise Exception("DOCX generation not supported")
        
        try:
            doc = Document()
            doc.add_heading(title, 0)
            
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            raise Exception("DOCX generation failed")
    
    def _detect_languages_from_url(self, repo_url):
        """Detect likely programming languages from repository URL without LLM"""
        default_langs = os.getenv("DEFAULT_LANGUAGES", ".js,.html,.css,.jsx,.ts,.tsx").split(',')
        return set([lang.strip() for lang in default_langs])
    
    def _get_frontend_extensions(self, architecture_content):
        """Get frontend file extensions without LLM"""
        frontend_exts = os.getenv("FRONTEND_EXTENSIONS", ".js,.jsx,.ts,.tsx,.vue,.html,.css").split(',')
        return [ext.strip() for ext in frontend_exts]
    
    def _get_programming_extensions(self, repo_url):
        """Get programming extensions without LLM"""
        prog_exts = os.getenv("PROGRAMMING_EXTENSIONS", ".py,.js,.jsx,.ts,.tsx,.java,.go,.rs,.php,.rb,.cs,.html,.css").split(',')
        return [ext.strip() for ext in prog_exts]

    def clone_and_analyze_repo(self, repo_url):
        """Clone and analyze GitHub repository"""
        temp_dir = tempfile.mkdtemp()
        
        try:
            clone_commands = [
                ['git', 'clone', '--depth', '1', repo_url, temp_dir],
                ['git', 'clone', repo_url, temp_dir]
            ]
            
            clone_success = False
            last_error = None
            
            for cmd in clone_commands:
                try:
                    timeout_val = int(os.getenv("GIT_CLONE_TIMEOUT", "60"))
                    subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=timeout_val)
                    clone_success = True
                    break
                except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
                    last_error = e
                    continue
            
            if not clone_success:
                # Return basic analysis without cloning
                return {
                    'files': {},
                    'structure': ['Repository clone failed - analysis based on URL only'],
                    'languages': self._detect_languages_from_url(repo_url),
                    'frameworks': [],
                    'dependencies': {},
                    'clone_error': str(last_error)
                }
            
            repo_analysis = {
                'files': {},
                'structure': [],
                'languages': set(),
                'frameworks': [],
                'dependencies': {}
            }
            
            for root, dirs, files in os.walk(temp_dir):
                excluded_dirs = os.getenv("EXCLUDED_DIRS", "node_modules,__pycache__,.git,.vscode").split(',')
                dirs[:] = [d for d in dirs if not d.startswith('.') and d not in [ex.strip() for ex in excluded_dirs]]
                
                for file in files:
                    if file.startswith('.'):
                        continue
                        
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, temp_dir)
                    file_ext = os.path.splitext(file)[1].lower()
                    
                    repo_analysis['structure'].append(rel_path)
                    
                    prog_exts = self._get_programming_extensions(repo_url)
                    if file_ext in prog_exts:
                        repo_analysis['languages'].add(file_ext)
            
            return repo_analysis
            
        except Exception as e:
            # Return fallback analysis
            return {
                'files': {},
                'structure': ['Repository analysis unavailable'],
                'languages': self._detect_languages_from_url(repo_url),
                'frameworks': [],
                'dependencies': {},
                'analysis_error': str(e)
            }
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def call_openai_api(self, prompt, max_retries=3):
        """Make API call to OpenAI with retry logic"""
        if not self.api_key:
            raise Exception("OPENROUTER_API_KEY not configured")
            
        for attempt in range(max_retries):
            try:
                data = {
                    "messages": [{"role": "user", "content": prompt}],
                    "model": self.model,
                    "temperature": self.temperature,
                    "max_tokens": int(os.getenv("MAX_TOKENS", "4500"))
                }
                
                headers = {
                    'Authorization': f'Bearer {self.api_key}',
                    'Content-Type': 'application/json'
                }
                
                req = urllib.request.Request(
                    self.api_url,
                    data=json.dumps(data).encode('utf-8'),
                    headers=headers
                )
                
                with urllib.request.urlopen(req) as response:
                    result = json.loads(response.read().decode('utf-8'))
                    return result['choices'][0]['message']['content']
                    
            except urllib.error.HTTPError as e:
                error_body = e.read().decode('utf-8')
                if e.code == 429:
                    if attempt < max_retries - 1:
                        base_wait = int(os.getenv("RATE_LIMIT_BASE_WAIT", "60"))
                        wait_time = base_wait * (2 ** attempt)
                        time.sleep(wait_time)
                        continue
                    else:
                        raise Exception("Rate limit exceeded. Please wait and try again.")
                else:
                    raise Exception(f"API Error {e.code}: {error_body}")
            except Exception as e:
                if attempt < max_retries - 1:
                    time.sleep(1)
                    continue
                raise Exception(f"API call failed: {str(e)}")
        
        raise Exception("Max retries exceeded")

    def analyze_with_openai(self, repo_url, repo_analysis, architecture_content, prd_content=None):
        """Generate comprehensive analysis using OpenAI API"""
        try:
            # Truncate architecture content if too long
            max_arch_chars = int(os.getenv("MAX_ARCH_CHARS", "3000"))
            if len(architecture_content) > max_arch_chars:
                architecture_content = architecture_content[:max_arch_chars] + "\n[Content truncated due to length]"
            
            # Extract key repo info
            frontend_exts = os.getenv("FRONTEND_EXTENSIONS", ".js,.jsx,.ts,.tsx,.vue,.html,.css").split(',')
            frontend_files = [f for f in repo_analysis['structure'][:5] if any(f.endswith(ext.strip()) for ext in frontend_exts)]
            languages = list(repo_analysis['languages'])[:3]
            
            prompt = f"""Analyze this frontend repo and architecture document to provide comprehensive tech stack justification and alternatives:

REPO: {repo_url}
LANGUAGES: {', '.join(languages)}
FRONTEND FILES: {', '.join(frontend_files)}

ARCHITECTURE:
{architecture_content}

PRD CONTENT:
{prd_content or 'Not provided'}

Generate output with:

# PROJECT SUMMARY
**Repository**: {repo_url}
[Generate a concise project summary based on the repository analysis and architecture document]

# ARCHITECTURE DIAGRAM
[Create an ASCII-based system architecture diagram showing the flow between frontend, backend, database, and external services. Use boxes, arrows, and clear labels]

# TECH STACK JUSTIFICATION
[For each technology mentioned in the architecture document, provide detailed justification why it's suitable for this project, including pros/cons and fit with requirements]

# ALTERNATIVE TECH STACKS
## Backend Alternatives:
- Option 1: [Framework/Language] - Pros, Cons, Performance, Scalability
- Option 2: [Framework/Language] - Pros, Cons, Performance, Scalability
- Option 3: [Framework/Language] - Pros, Cons, Performance, Scalability

## Database Alternatives:
- Option 1: [Database Type] - Use cases, Performance, Cost
- Option 2: [Database Type] - Use cases, Performance, Cost
- Option 3: [Database Type] - Use cases, Performance, Cost

# DATABASE SCHEMA DESIGN
[Based on architecture requirements, design database tables with fields, data types, relationships, and indexes]

# RECOMMENDED API ENDPOINTS
[Based on the architecture requirements, suggest complete API specification with:
- HTTP methods (GET, POST, PUT, DELETE)
- Full endpoint paths
- Request/response body formats
- Authentication requirements
- Input validation rules
- Error response formats]

## Input Fields for Each Endpoint:
[For each API endpoint, specify:
- Required input fields with data types
- Optional input fields with defaults
- Validation rules (min/max length, format, etc.)
- Example request payloads
- Field descriptions and purposes]

# DETAILED PROJECT CONSTRUCTION GUIDE
[Use the BEST ALTERNATIVE tech stack from the alternatives analysis above. If architecture document tech stack is clearly superior, use that instead. Clearly state which tech stack combination you're using for this guide.]

## Phase 1: Environment Setup
1. Install [SPECIFIC BACKEND FRAMEWORK] development environment
2. Setup [SPECIFIC DATABASE] server and tools
3. Initialize Git repository with proper .gitignore
4. Create project folder structure for chosen tech stack
5. Setup package managers and dependency files
6. Configure environment variables template

## Phase 2: Backend Development
1. Install [SPECIFIC FRAMEWORK] and create initial project structure
2. Configure [SPECIFIC DATABASE] connection with credentials
3. Create database models/entities based on schema design above
4. Create input validation schemas for all API endpoints
5. Implement each API endpoint with detailed input/output handling
6. Setup JWT/OAuth authentication with input validation
7. Add middleware for logging, CORS, rate limiting, input sanitization
8. Create comprehensive test suites for all endpoints and input validation

## Phase 3: Database Implementation
1. Setup [SPECIFIC DATABASE TYPE] server (local/cloud)
2. Run database migrations to create schema from design above
3. Create seed scripts with sample data
4. Add database indexes for performance optimization
5. Setup automated backup procedures
6. Configure connection pooling and optimization

## Phase 4: Integration
1. Connect existing frontend to new backend APIs
2. Update frontend API calls to match new endpoints
3. Implement error handling and loading states
4. Add authentication flow integration
5. Test all frontend-backend data flows
6. Optimize API response times and caching

Provide SPECIFIC commands, code snippets, and configuration examples for the CHOSEN tech stack in each step."""
            
            return self.call_openai_api(prompt)
            
        except Exception as e:
            return f"Error generating analysis: {str(e)}"

# Initialize agent
agent = ImpactAnalysisAgent()

# Routes
@app.get("/", response_class=HTMLResponse)
async def read_root():
    """Serve the main HTML page"""
    try:
        with open("static/index.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>Frontend files not found. Please ensure static files are deployed.</h1>")

@app.post("/upload-file", response_model=FileUploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """Handle file upload and text extraction"""
    try:
        file_data = await file.read()
        extracted_text = agent.extract_text_from_file(file_data, file.filename)
        
        return FileUploadResponse(
            success=True,
            extracted_text=extracted_text,
            filename=file.filename
        )
    except Exception as e:
        return FileUploadResponse(
            success=False,
            error=str(e)
        )

@app.post("/analyze", response_model=AnalysisResponse)
async def analyze_project(request: AnalysisRequest):
    """Handle analysis requests"""
    try:
        repo_analysis = agent.clone_and_analyze_repo(request.repo_url)
        analysis = agent.analyze_with_openai(
            request.repo_url, 
            repo_analysis, 
            request.architecture_content, 
            request.prd_content
        )
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        os.makedirs(GENERATED_FILES_DIR, exist_ok=True)
        
        documents = {
            'prompt': analysis,
            'architecture': request.architecture_content,
            'prd': request.prd_content or 'No PRD provided',
            'repository_url': request.repo_url
        }
        
        with open(f'{GENERATED_FILES_DIR}/documents_{timestamp}.json', 'w', encoding='utf-8') as f:
            json.dump(documents, f, indent=2)
        
        return AnalysisResponse(
            success=True,
            analysis=analysis,
            document_id=timestamp,
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        return AnalysisResponse(
            success=False,
            error=str(e)
        )

@app.get("/download/{document_id}/{doc_type}/{format_type}")
async def download_document(document_id: str, doc_type: str, format_type: str):
    """Handle document downloads"""
    try:
        json_path = f'{GENERATED_FILES_DIR}/documents_{document_id}.json'
        
        if not os.path.exists(json_path):
            raise HTTPException(status_code=404, detail="Document not found")
        
        with open(json_path, 'r', encoding='utf-8') as f:
            documents = json.load(f)
        
        if doc_type not in documents:
            raise HTTPException(status_code=404, detail="Document type not found")
        
        content = documents[doc_type]
        filename = f"{doc_type}_{document_id}"
        
        if format_type == 'txt':
            return StreamingResponse(
                BytesIO(content.encode('utf-8')),
                media_type='text/plain',
                headers={"Content-Disposition": f"attachment; filename={filename}.txt"}
            )
        elif format_type == 'pdf':
            if not PDF_SUPPORT:
                raise HTTPException(status_code=400, detail="PDF not supported")
            pdf_data = agent.generate_pdf(content, f"{doc_type.title()} - {document_id}")
            return StreamingResponse(
                BytesIO(pdf_data),
                media_type='application/pdf',
                headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
            )
        else:
            raise HTTPException(status_code=400, detail="Invalid format")
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def main():
    """Main function to start the server"""
    try:
        # Handle Render's PORT environment variable properly
        port_env = os.getenv("PORT")
        
        # Use Render's assigned port or fallback to 10000
        if port_env and port_env.isdigit():
            port_to_use = int(port_env)
        else:
            port_to_use = 10000
            
        host = "0.0.0.0"
        
        os.makedirs(GENERATED_FILES_DIR, exist_ok=True)
        
        print(f"Starting server on {host}:{port_to_use}")
        uvicorn.run(app, host=host, port=port_to_use, log_level="info")
    except Exception as e:
        print(f"Failed to start server: {e}")
        raise

if __name__ == "__main__":
    main()